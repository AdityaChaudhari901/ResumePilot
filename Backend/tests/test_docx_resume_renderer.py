from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

from docx import Document

from app.services.docx_resume_renderer import render_tailored_resume_docx
from app.services.job_parser import parse_job_profile
from app.services.matcher import match_resume_to_job
from app.services.report_generator import generate_report
from app.services.resume_parser import parse_resume_profile


def test_docx_renderer_builds_evidence_backed_resume_and_excludes_missing_skills():
    resume = parse_resume_profile(
        """A&B Dev
a.dev@example.com
https://github.com/A_B

Skills
Python, FastAPI, PostgreSQL

Projects
Built FastAPI API services with PostgreSQL and improved reliability by 40%.
""",
        resume_id=1,
    )
    job = parse_job_profile(
        """Role: Platform Engineer
Company: Example Labs

Requirements:
- Required Python experience.
- Required FastAPI experience.
- Required Kubernetes experience.
""",
        job_id=1,
    )
    match = match_resume_to_job(resume, job)
    report = generate_report(analysis_id=1, resume=resume, job=job, match=match)

    docx_bytes = render_tailored_resume_docx(report=report, resume=resume, job=job)

    assert docx_bytes.startswith(b"PK")
    with ZipFile(BytesIO(docx_bytes)) as archive:
        assert "word/document.xml" in archive.namelist()

    document = Document(BytesIO(docx_bytes))
    exported_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading")
    ]

    assert document.core_properties.author == "ResumePilot"
    assert document.core_properties.last_modified_by == "ResumePilot"
    assert "A&B Dev" in exported_text
    assert "github.com/A_B" in exported_text
    assert "Professional Summary" in headings
    assert "Technical Skills" in headings
    assert "Evidence-Backed Tailored Highlights" in headings
    assert "Built FastAPI API services with PostgreSQL" in exported_text
    assert "Python" in exported_text
    assert "FastAPI" in exported_text
    assert "Kubernetes" not in exported_text
