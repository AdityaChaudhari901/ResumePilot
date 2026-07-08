from __future__ import annotations

from collections import defaultdict
from io import BytesIO
from urllib.parse import urlparse

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.schemas.job import JobProfile
from app.schemas.report import ApplicationReport, TailoredBullet
from app.schemas.resume import CandidateProfile, ResumeFact, ResumeProfile, ResumeSkill

MAX_SUMMARY_FACTS = 2
MAX_TAILORED_BULLETS = 5
MAX_EXPERIENCE_FACTS = 7
MAX_PROJECT_FACTS = 6

SECTION_LABELS = {
    "programming_language": "Languages",
    "backend_framework": "Backend",
    "frontend_framework": "Frontend",
    "database": "Databases",
    "cloud_devops": "DevOps/Tools",
    "ai_ml": "AI/ML",
    "data_tool": "Data Tools",
    "testing": "Testing",
    "security": "Security",
    "soft_skill": "Soft Skills",
    "domain": "Domain",
    "other": "Other",
}

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
BLACK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x4B, 0x55, 0x63)


def render_tailored_resume_docx(
    *,
    report: ApplicationReport,
    resume: ResumeProfile,
    job: JobProfile,
) -> bytes:
    """Render an evidence-backed tailored resume as a DOCX byte stream."""

    document = Document()
    _configure_document(document)
    _render_header(document, resume.candidate)
    _render_summary_section(document, resume, report, job)
    _render_skills_section(document, resume.skills, report)
    _render_tailored_section(document, report.tailored_bullets)
    _render_fact_section(
        document,
        title="Professional Experience",
        facts=resume.experience,
        max_items=MAX_EXPERIENCE_FACTS,
    )
    _render_fact_section(
        document,
        title="Projects",
        facts=resume.projects,
        max_items=MAX_PROJECT_FACTS,
    )
    _render_plain_fact_section(document, "Education", resume.education)
    _render_plain_fact_section(document, "Certifications", resume.certifications)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _configure_document(document: DocxDocument) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    core_properties = document.core_properties
    core_properties.author = "ResumePilot"
    core_properties.last_modified_by = "ResumePilot"
    core_properties.title = "Evidence-backed tailored resume"
    core_properties.subject = "ResumePilot resume export"
    core_properties.comments = (
        "Generated from validated ResumeProfile, JobProfile, and ApplicationReport data."
    )
    core_properties.keywords = "resume,resumepilot,evidence-backed"

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, font_size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 12, BLUE),
        ("Heading 3", 11, DARK_BLUE),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True

    bullet_style = styles["List Bullet"]
    bullet_style.font.name = "Calibri"
    bullet_style.font.size = Pt(10)
    bullet_style.paragraph_format.left_indent = Inches(0.25)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.13)
    bullet_style.paragraph_format.space_after = Pt(2)
    bullet_style.paragraph_format.line_spacing = 1.1


def _render_header(document: DocxDocument, candidate: CandidateProfile) -> None:
    name = candidate.name or "Candidate"
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(name)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = BLACK

    contact_items = _contact_items(candidate)
    if contact_items:
        contact = document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.paragraph_format.space_after = Pt(6)
        contact_run = contact.add_run(" | ".join(contact_items))
        contact_run.font.name = "Calibri"
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = MUTED


def _contact_items(candidate: CandidateProfile) -> list[str]:
    items: list[str] = []
    if candidate.location:
        items.append(candidate.location)
    if candidate.phone:
        items.append(candidate.phone)
    if candidate.email:
        items.append(str(candidate.email))
    for link in candidate.links:
        items.append(_display_url(str(link).rstrip("/")))
    return items


def _render_summary_section(
    document: DocxDocument,
    resume: ResumeProfile,
    report: ApplicationReport,
    job: JobProfile,
) -> None:
    summary_facts = [
        fact.text.strip()
        for fact in resume.facts
        if fact.section == "summary" and fact.text.strip()
    ][:MAX_SUMMARY_FACTS]
    if summary_facts:
        summary = " ".join(summary_facts)
    else:
        top_matches = [skill.skill for skill in report.matched_skills[:4]]
        role_context = _role_context(job)
        if top_matches:
            summary = (
                f"Evidence-backed technical resume tailored for {role_context}, "
                f"with supported experience in {_human_list(top_matches)}."
            )
        else:
            summary = f"Evidence-backed technical resume tailored for {role_context}."
    _render_paragraph_section(document, "Professional Summary", summary)


def _render_skills_section(
    document: DocxDocument,
    skills: list[ResumeSkill],
    report: ApplicationReport,
) -> None:
    evidence_backed_skills = [skill for skill in skills if skill.evidence_ids]
    if not evidence_backed_skills:
        return

    matched_names = {_normalize_skill(skill.skill) for skill in report.matched_skills}
    grouped: dict[str, list[ResumeSkill]] = defaultdict(list)
    for skill in sorted(
        evidence_backed_skills,
        key=lambda item: (_normalize_skill(item.name) not in matched_names, item.name.lower()),
    ):
        grouped[skill.category.value].append(skill)

    _add_section_heading(document, "Technical Skills")
    for category, label in SECTION_LABELS.items():
        category_skills = grouped.get(category, [])
        if not category_skills:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.name = "Calibri"
        label_run.font.size = Pt(10)
        paragraph.add_run(", ".join(skill.name for skill in category_skills))


def _render_tailored_section(
    document: DocxDocument,
    bullets: list[TailoredBullet],
) -> None:
    safe_bullets = [
        bullet.bullet for bullet in bullets if bullet.evidence_ids and not bullet.unsupported_claims
    ][:MAX_TAILORED_BULLETS]
    _render_bullet_section(document, "Evidence-Backed Tailored Highlights", safe_bullets)


def _render_fact_section(
    document: DocxDocument,
    *,
    title: str,
    facts: list[ResumeFact],
    max_items: int,
) -> None:
    _render_bullet_section(document, title, [fact.text for fact in facts[:max_items]])


def _render_plain_fact_section(
    document: DocxDocument,
    title: str,
    facts: list[ResumeFact],
) -> None:
    clean_facts = [fact.text.strip() for fact in facts if fact.text.strip()]
    if not clean_facts:
        return
    _add_section_heading(document, title)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(" | ".join(clean_facts))


def _render_paragraph_section(
    document: DocxDocument,
    title: str,
    body: str,
) -> None:
    clean_body = body.strip()
    if not clean_body:
        return
    _add_section_heading(document, title)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(clean_body)


def _render_bullet_section(
    document: DocxDocument,
    title: str,
    bullets: list[str],
) -> None:
    clean_bullets = [bullet.strip() for bullet in bullets if bullet.strip()]
    if not clean_bullets:
        return
    _add_section_heading(document, title)
    for bullet in clean_bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(bullet)


def _add_section_heading(document: DocxDocument, title: str) -> None:
    paragraph = document.add_heading(title, level=2)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)


def _display_url(value: str) -> str:
    parsed = urlparse(value)
    if parsed.netloc:
        return f"{parsed.netloc}{parsed.path}".rstrip("/")
    return value


def _role_context(job: JobProfile) -> str:
    role = job.role_title or "the target role"
    if job.company:
        return f"{role} at {job.company}"
    return role


def _human_list(values: list[str]) -> str:
    if not values:
        return "relevant evidence-backed skills"
    if len(values) == 1:
        return values[0]
    return f"{', '.join(values[:-1])}, and {values[-1]}"


def _normalize_skill(value: str) -> str:
    return " ".join(value.casefold().split())
